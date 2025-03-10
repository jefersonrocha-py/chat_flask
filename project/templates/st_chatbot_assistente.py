import os
import requests
import logging
import pandas as pd
import streamlit as st
from io import StringIO
from docx import Document
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from langchain.schema import Document as LangChainDocument
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.chat_models import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
import json
import hashlib  # [Melhoria: Otimização de Desempenho] Para gerar hash do arquivo

# Configuração de logging
logging.basicConfig(
    filename="chatbot_assistente_errors.log",
    level=logging.ERROR,
    format="%(asctime)s - %(levelname)s - %(message)s"
)

# Função para exibir mensagens de erro amigáveis
def show_error(message):
    st.error(f"😕 Oops! Algo deu errado: {message}")
    logging.error(message)

# Carregar variáveis de ambiente
try:
    load_dotenv()
except Exception as e:
    show_error(f"Não foi possível carregar variáveis de ambiente: {e}")

# Configurações do Ollama (definir antes de qualquer uso)
OLLAMA_SERVER_URL = "http://localhost:11434"

# Verificar conexão com o servidor backend
def check_backend_connection():
    OLLAMA_SERVER_URL = "http://localhost:11434"
    try:
        response = requests.get(OLLAMA_SERVER_URL, timeout=5)
        if response.status_code == 200:
            return True
        else:
            st.error(f"Erro ao conectar ao servidor backend: Status Code {response.status_code}")
    except requests.exceptions.Timeout:
        st.error(f"Timeout ao tentar conectar ao servidor backend em {OLLAMA_SERVER_URL}.")
    except Exception as e:
        st.error(f"Erro ao conectar ao servidor backend: {e}")
    return False

# Verificar conexão antes de prosseguir
backend_connected = check_backend_connection()

if not backend_connected:
    st.warning("⚠️ O servidor backend não está disponível. Algumas funcionalidades podem estar limitadas.")

# Configurar o Ollama como provedor de embeddings e chat
try:
    embeddings = OllamaEmbeddings(base_url=OLLAMA_SERVER_URL, model="llama3.2:latest")
    chat_model = ChatOllama(base_url=OLLAMA_SERVER_URL, model="llama3.2:latest")
except Exception as e:
    show_error(f"Erro ao configurar embeddings ou chat model: {e}")
    st.stop()

# Criar diretório para armazenar arquivos
DOCUMENTS_DIR = "utils/uploads/files"
try:
    os.makedirs(DOCUMENTS_DIR, exist_ok=True)
    # [Melhoria: Segurança] Definir permissões restritas para o diretório
    os.chmod(DOCUMENTS_DIR, 0o700)  # Apenas o dono pode acessar
except Exception as e:
    show_error(f"Erro ao criar diretório: {e}")
    st.stop()

# Funções para processar diferentes tipos de arquivos
def process_csv(file_path):
    """Processa arquivos CSV."""
    try:
        df = pd.read_csv(file_path)
        return [{"page_content": row.to_string()} for _, row in df.iterrows()]
    except Exception as e:
        show_error(f"Erro ao processar CSV: {e}")
        return None

def process_xlsx(file_path):
    """Processa arquivos Excel."""
    try:
        df = pd.read_excel(file_path)
        return [{"page_content": row.to_string()} for _, row in df.iterrows()]
    except Exception as e:
        show_error(f"Erro ao processar Excel: {e}")
        return None

def process_docx(file_path):
    """Processa arquivos Word."""
    try:
        doc = Document(file_path)
        full_text = [p.text for p in doc.paragraphs if p.text.strip()]
        return [{"page_content": text} for text in full_text]
    except Exception as e:
        show_error(f"Erro ao processar DOCX: {e}")
        return None

def process_pdf(file_path):
    """Processa arquivos PDF."""
    try:
        reader = PdfReader(file_path)
        pages = [page.extract_text() for page in reader.pages if page.extract_text()]
        return [{"page_content": text} for text in pages]
    except Exception as e:
        show_error(f"Erro ao processar PDF: {e}")
        return None

def process_txt(file_path):
    """Processa arquivos TXT."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return [{"page_content": text}]
    except Exception as e:
        show_error(f"Erro ao processar TXT: {e}")
        return None

# [Melhoria: Segurança] Função para validar o tipo de arquivo
def validate_file(uploaded_file):
    allowed_types = ['csv', 'xml', 'xls', 'xlsx', 'docx', 'pdf', 'txt']
    ext = uploaded_file.name.split('.')[-1].lower()
    if ext not in allowed_types:
        raise ValueError("Tipo de arquivo não suportado. Use .csv, .xml, .xls, .xlsx, .docx, .pdf ou .txt.")

# [Melhoria: Otimização de Desempenho] Função para gerar hash do arquivo
def get_file_hash(file_path):
    try:
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    except Exception as e:
        show_error(f"Erro ao gerar hash do arquivo: {e}")
        return None

@st.cache_data
def load_data(file_path, _file_hash):  # [Melhoria: Otimização de Desempenho] Adicionado _file_hash para cache dinâmico
    """Carrega o arquivo e gera o retriever para busca."""
    try:
        file_type = file_path.split('.')[-1].lower()
        if file_type == 'csv':
            documents_raw = process_csv(file_path)
        elif file_type in ['xlsx', 'xls', 'xml']:
            documents_raw = process_xlsx(file_path)
        elif file_type == 'docx':
            documents_raw = process_docx(file_path)
        elif file_type == 'pdf':
            documents_raw = process_pdf(file_path)
        elif file_type == 'txt':
            documents_raw = process_txt(file_path)
        else:
            show_error("Formato de arquivo não suportado. Use .csv, .xlsx, .docx, .pdf ou .txt.")
            return None
        if not documents_raw:
            show_error("Nenhum dado foi processado do arquivo.")
            return None
        documents = [LangChainDocument(page_content=doc["page_content"]) for doc in documents_raw]
        vectorstore = FAISS.from_documents(documents, embeddings)
        return vectorstore.as_retriever()
    except Exception as e:
        show_error(f"Erro ao carregar dados: {e}")
        return None

# Função para verificar arquivos no diretório
def check_files_in_directory():
    """Verifica se existem arquivos no diretório e retorna o caminho do primeiro arquivo encontrado."""
    try:
        files = os.listdir(DOCUMENTS_DIR)
        if files:
            return os.path.join(DOCUMENTS_DIR, files[0])
        return None
    except Exception as e:
        show_error(f"Erro ao verificar diretório: {e}")
        return None

# CSS customizado para estilização (incluindo o footer fixo e ajuste do chat input)
st.markdown(
    """
    <style>
    /* Estilo geral */
    .css-18e3th9 {
        padding-top: 1rem;
    }
    .css-1d391kg {
        padding: 0;
    }
    .chat-title {
        font-size: 32px;
        font-weight: bold;
        text-align: center;
        margin-bottom: 10px;
        color: #4CAF50;
    }
    .avatar-container {
        display: flex;
        align-items: center;
        justify-content: center;
        margin-bottom: 20px;
    }
    .avatar-initial {
        width: 60px;
        height: 60px;
        border-radius: 50%;
        background-color: #4CAF50;
        color: white;
        font-size: 24px;
        font-weight: bold;
        display: flex;
        align-items: center;
        justify-content: center;
        margin-right: 0.5rem;
    }

    /* Botão Voltar */
    .back-button {
        background-color: #2196F3; /* Azul */
        color: white;
        padding: 0.6em 1em;
        border: none;
        border-radius: 20px; /* Arredondado */
        cursor: pointer;
        text-decoration: none;
        font-size: 16px;
        font-weight: bold;
        box-shadow: none; /* Sem sombra */
        outline: none; /* Remove o contorno branco */
        transition: background-color 0.3s ease; /* Transição suave */
    }
    /* Adaptação para tema Light */
    [data-testid="stSidebar"] .back-button {
        background-color: #2196F3; /* Azul padrão */
    }
    /* Adaptação para tema Dark */
    [data-testid="stSidebar"][class*="dark"] .back-button {
        background-color: #1E88E5; /* Azul mais escuro para o tema Dark */
    }

    /* Rodapé da sidebar */
    .sidebar-footer {
        position: fixed;
        bottom: 0;
        width: 250px; /* Largura da sidebar */
        background-color: transparent; /* Fundo transparente para adaptar ao tema */
        padding: 10px 0;
        text-align: center;
        box-shadow: none; /* Sem sombra */
    }

    /* Ajusta o input do chat para não sobrepor o footer */
    [data-testid="stChatInput"] {
        margin-bottom: 90px !important; /* Aumente/diminua conforme necessário */
    }

    /* Footer fixo */
    .custom-footer {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background-color: #0E1117; /* Ajuste para a cor desejada */
        padding: 1rem 5%;
        display: flex;
        justify-content: space-between;
        align-items: center;
        z-index: 9999;
        border-top: 1px solid #2f2f2f;
        color: #fff;
    }

    .footer-content {
        max-width: 1200px;
        margin: 0 auto;
        display: flex;
        justify-content: space-between;
        align-items: center;
        width: 100%;
    }

    .footer-logo img {
        height: 40px;
        filter: brightness(0) invert(var(--logo-invert, 0));
        transition: transform 0.3s;
    }

    .footer-center {
        flex-grow: 1;
        text-align: center;
        font-size: 0.95rem;
        margin: 0 1.5rem;
    }

    .footer-neon {
        animation: neonPulse 1.5s infinite alternate;
        white-space: nowrap;
    }

    @keyframes neonPulse {
        from { text-shadow: 0 0 5px rgba(8,255,184,0.3); }
        to { text-shadow: 0 0 15px rgba(8,255,184,0.5); }
    }

    @media (max-width: 768px) {
        .footer-content {
            flex-direction: column;
            gap: 1rem;
        }
        .footer-center {
            order: 2;
            margin: 0;
        }
        .footer-logo {
            order: 1;
        }
        .footer-neon {
            order: 3;
        }
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# Obter o username dos parâmetros da URL
query_params = st.query_params
user_logged = query_params.get("username", ["Usuário"])[0]

def chatbot_assistente_page():
    st.markdown(
        """
        <div class="chat-title">FlowMind AI 🤖</div>
        """,
        unsafe_allow_html=True
    )

    # Verifica se há um arquivo disponível no diretório
    initial_file = check_files_in_directory()
    retriever = None
    if initial_file:
        # Verifica se já existe um retriever carregado na sessão
        if "retriever" not in st.session_state or st.session_state["current_file"] != initial_file:
            st.info(f"📄 Arquivo detectado no diretório: {os.path.basename(initial_file)}. Carregando...")
            file_hash = get_file_hash(initial_file)
            retriever = load_data(initial_file, _file_hash=file_hash)
            st.session_state["retriever"] = retriever
            st.session_state["current_file"] = initial_file
        else:
            retriever = st.session_state["retriever"]

    with st.sidebar:
        initial = user_logged[0].upper() if user_logged else "U"
        st.markdown(
            f"""
            <div class="avatar-container">
                <div class="avatar-initial">{initial}</div>
            </div>
            """,
            unsafe_allow_html=True
        )

        st.header("📁 Upload dos Files")
        # [Melhoria: Permitir Múltiplos Uploads de Arquivos] Alterado para aceitar múltiplos arquivos
        uploaded_files = st.file_uploader(
            "Envie sua base de conhecimento (.csv, .xml, .xls, .xlsx, .docx, .pdf, .txt)",
            type=['csv', 'xml', 'xls', 'xlsx', 'docx', 'pdf', 'txt'],
            accept_multiple_files=True
        )
        if uploaded_files:
            for uploaded_file in uploaded_files:
                try:
                    # [Melhoria: Segurança] Validar o arquivo antes de processar
                    validate_file(uploaded_file)
                    file_path = os.path.join(DOCUMENTS_DIR, uploaded_file.name)
                    with open(file_path, 'wb') as f:
                        f.write(uploaded_file.getbuffer())
                    # [Melhoria: Interface do Usuário] Feedback positivo com nome do arquivo
                    st.success(f"✅ Arquivo {uploaded_file.name} carregado com sucesso!")
                    # [Melhoria: Interface do Usuário] Indicador de progresso durante o processamento
                    with st.spinner(f"Processando {uploaded_file.name}..."):
                        file_hash = get_file_hash(file_path)
                        retriever = load_data(file_path, _file_hash=file_hash)
                    st.session_state["retriever"] = retriever
                    st.session_state["current_file"] = file_path
                except ValueError as ve:
                    show_error(str(ve))
                except Exception as e:
                    show_error(f"Erro ao salvar ou processar arquivo: {e}")

        # Botão "Voltar" no rodapé da sidebar
        FLASK_ROUTE = "http://localhost:5000/mode_selection"
        st.markdown(
            f"""
            <div class="sidebar-footer">
                <a href="{FLASK_ROUTE}" target="_self" style="text-decoration: none;">
                    <button class="back-button">🔙 Voltar</button>
                </a>
            </div>
            """,
            unsafe_allow_html=True
        )

    # Se houver retriever, exibe a interface do chat
    if retriever:
        # [Melhoria: Melhoria no Modelo - Ajuste no Prompt] Template ajustado para respostas mais precisas
        rag_template = """
        Você é um assistente especializado. 
        Seu trabalho é fornecer respostas claras, concisas e precisas com base nas informações fornecidas no arquivo.
        Contexto: {context}
        Pergunta do cliente: {question}
        """
        prompt = ChatPromptTemplate.from_template(rag_template)
        chain = (
            {"context": retriever, "question": RunnablePassthrough()}
            | prompt
            | chat_model
        )
        if "messages" not in st.session_state:
            st.session_state.messages = []

        # Exibe o histórico de mensagens
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        # Entrada do usuário
        if user_input := st.chat_input("Você:"):
            st.session_state.messages.append({"role": "user", "content": user_input})
            with st.chat_message("user"):
                st.markdown(user_input)
            try:
                response_stream = chain.stream({"text": user_input})
                full_response = ""
                response_container = st.chat_message("assistant")
                response_text = response_container.empty()
                for partial_response in response_stream:
                    full_response += str(partial_response.content)
                    response_text.markdown(full_response + "▌")
                st.session_state.messages.append({"role": "assistant", "content": full_response})
                # [Melhoria: Salvar Histórico Persistente em Arquivo JSON] Salva o histórico após cada resposta
                save_history()
            except Exception as e:
                show_error(f"Erro ao gerar resposta: {e}")
    else:
        st.info("📂 Use a aba à esquerda para carregar um arquivo e começar.")

    # Footer fixo (exibido sempre)
    st.markdown("""
        <div class="custom-footer">
            <div class="footer-content">
                <div class="footer-logo">
                    <img src="https://etheriumtech.com.br/wp-content/uploads/2024/04/LOGO-BRANCO.png" alt="Logo">
                </div>
                <div class="footer-center">
                    FlowMind AI © 2025 - Todos os direitos reservados
                </div>
                <div class="footer-neon">
                    💡🔗 Powered by Jeferson Rocha
                </div>
            </div>
        </div>
        """, unsafe_allow_html=True)

# [Melhoria: Salvar Histórico Persistente em Arquivo JSON] Função para salvar o histórico em JSON
def save_history():
    try:
        with open("chat_assistente_history.json", "w") as f:
            json.dump(st.session_state.messages, f)
    except Exception as e:
        show_error(f"Erro ao salvar histórico: {e}")

# Executa a página do chatbot
try:
    chatbot_assistente_page()
except Exception as e:
    show_error(f"Erro crítico na aplicação: {e}")
